
import re
import os
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches, RGBColor  # <--- Añadido RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


# --- FUNCIONES DE APOYO ---

def aplicar_estilo_calibri(run, tamano, negrita=False, color_rgb=None):
    """
    Aplica estilo Calibri, tamaño, negrita y color opcional.
    """
    run.font.name = 'Calibri'
    run.font.size = Pt(tamano)
    run.bold = negrita

    if color_rgb:
        run.font.color.rgb = color_rgb  # <--- Aplicación del color

    r = run._element
    rFonts = r.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = r.makeelement(qn('w:rFonts'))
        r.insert(0, rFonts)
    rFonts.set(qn('w:ascii'), 'Calibri')
    rFonts.set(qn('w:hAnsi'), 'Calibri')


def extraer_datos_completos(ruta):
    if not os.path.exists(ruta):
        return None, None
    doc = Document(ruta)

    paciente_nombre = "No detectado"
    for p in doc.paragraphs:
        if "Paciente:" in p.text:
            paciente_nombre = p.text.split("Paciente:")[1].strip()
            break

    resumen_scores = {}
    if doc.tables:
        for row in doc.tables[0].rows[1:]:
            if len(row.cells) >= 2:
                resumen_scores[row.cells[0].text.strip()] = row.cells[1].text.strip().replace("*", "")

    datos_donantes = []
    perfil_paciente = {"Nombre": paciente_nombre, "EsPaciente": True,
                       "Loci": {"A": [], "B": [], "C": [], "DQ": [], "DRB1": []}}

    info_donantes_keys = ["A", "B", "C", "DQ", "DRB1"]
    from docx.oxml.table import CT_Tbl
    from docx.oxml.text.paragraph import CT_P
    from docx.table import Table
    from docx.text.paragraph import Paragraph

    ultimo_nombre = None
    for child in doc._body._body.iterchildren():
        if isinstance(child, CT_P):
            p = Paragraph(child, doc)
            if "Donante:" in p.text:
                ultimo_nombre = p.text.split("Donante:")[1].strip()

        elif isinstance(child, CT_Tbl) and ultimo_nombre:
            tabla = Table(child, doc)

            for row in tabla.rows[1:]:
                if len(row.cells) < 4: continue

                locus_raw = row.cells[0].text.strip().upper()
                val_pac = row.cells[1].text.strip()
                val_don = row.cells[2].text.strip()
                estado_texto = row.cells[3].text.strip().lower()

                key = "DQ" if "DQ" in locus_raw else "DRB1" if "DR" in locus_raw else locus_raw
                if key not in info_donantes_keys: continue

                # NUEVA LÓGICA DE ESTADOS
                es_compatible = estado_texto in ["compatible", "idéntico", "identico", "incluido"]
                es_incierto = "incierto" in estado_texto  # <--- Nueva condición

                # 1. Guardar alelo del donante
                info_donante = {"val": val_don, "bold": es_compatible, "es_incierto": es_incierto}

                # Buscamos si el donante ya existe en la lista, si no, lo creamos
                donante_actual = next((d for d in datos_donantes if d["Nombre"] == ultimo_nombre), None)
                if not donante_actual:
                    donante_actual = {"Nombre": ultimo_nombre, "EsPaciente": False,
                                      "Loci": {"A": [], "B": [], "C": [], "DQ": [], "DRB1": []}}
                    datos_donantes.append(donante_actual)

                donante_actual["Loci"][key].append(info_donante)

                # 2. Actualizar alelo del paciente
                lista_alelos_pac = perfil_paciente["Loci"][key]
                encontrado = False

                # Primero: Buscamos si el alelo ya existe para actualizar su estilo (negrita/rojo)
                for alelo_obj in lista_alelos_pac:
                    if alelo_obj["val"] == val_pac:
                        if es_compatible: alelo_obj["bold"] = True
                        if es_incierto: alelo_obj["es_incierto"] = True
                        encontrado = True
                        # No hacemos 'break' aquí para que si hay dos iguales, ambos se marquen
                        # como compatibles si corresponde en tablas futuras

                # SEGUNDO: Solo añadimos el alelo si NO lo hemos encontrado O si estamos
                # en la primera tabla y todavía no tenemos los 2 alelos del locus.
                # Esto permite la homocigosis (dos iguales) pero evita duplicados de otros donantes.

                es_primera_tabla = len(datos_donantes) <= 1
                limite_alelos_alcanzado = len(lista_alelos_pac) >= 2

                if val_pac and es_primera_tabla and not limite_alelos_alcanzado:
                    lista_alelos_pac.append({"val": val_pac, "bold": es_compatible, "es_incierto": es_incierto})

            ultimo_nombre = None

    return [perfil_paciente] + datos_donantes, resumen_scores


def generar_documento(lista_personas, resumen, ruta_salida):
    doc = Document()

    # --- CONFIGURACIÓN DE MÁRGENES AMPLIADOS ---
    # Margen de 2.54 cm (1 pulgada) en todos los lados para un look limpio
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Definición de tamaños
    T_TITULO_GRANDE = 20
    T_TEXTO = 11
    T_TABLA = 9

    nombre_paciente = lista_personas[0]['Nombre'].upper()

    # --- TÍTULO MÁS GRANDE Y CENTRADO ---
    tit = doc.add_paragraph()
    tit.alignment = WD_ALIGN_PARAGRAPH.CENTER
    aplicar_estilo_calibri(tit.add_run("LABORATORIO DE HISTOCOMPATIBILIDAD\nY CRIOPRESERVACION"), T_TITULO_GRANDE, True)

    doc.add_paragraph()  # Espacio

    # --- CUADRO DE DATOS COMPACTO ---
    head_tab = doc.add_table(rows=4, cols=2)
    head_tab.width = Inches(6.5)
    head_tab.alignment = WD_TABLE_ALIGNMENT.CENTER

    def llenar_celda(fila, col, tag, valor, bold_val=False):
        p = head_tab.cell(fila, col).paragraphs[0]
        p.paragraph_format.space_after = Pt(2)
        aplicar_estilo_calibri(p.add_run(tag), T_TEXTO, True)
        aplicar_estilo_calibri(p.add_run(valor), T_TEXTO, bold_val)

    llenar_celda(0, 0, "PACIENTE: ", nombre_paciente)
    llenar_celda(0, 1, "CÓDIGO: ", "HLA-2025")
    llenar_celda(1, 0, "FECHA: ", datetime.now().strftime('%d/%m/%Y'))
    llenar_celda(1, 1, "SERVICIO: ", "")
    llenar_celda(2, 0, "MÉDICO SOLICITANTE: ", "")
    llenar_celda(2, 1, "DIAGNÓSTICO: ", "")
    llenar_celda(3, 0, "INDICACIÓN: ", "")
    llenar_celda(3, 1, "MUESTRA: ", "SANGRE PERIFÉRICA")

    doc.add_paragraph()

    # --- SUBTÍTULO ---
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run_sub = sub.add_run("ESTUDIO FAMILIAR DE HISTOCOMPATIBILIDAD")
    aplicar_estilo_calibri(run_sub, 12, True)
    run_sub.underline = True

    # --- TABLA DE RESULTADOS (Lógica de compatibilidad) ---
    tab = doc.add_table(rows=1, cols=6)
    tab.style = 'Table Grid'
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER

    headers = ["Nombre", "Locus A", "Locus B", "Locus C", "Locus DQ", "Locus DRB1"]
    for i, h in enumerate(headers):
        p = tab.rows[0].cells[i].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        aplicar_estilo_calibri(p.add_run(h), 9, True)

    for pers in lista_personas:
        row = tab.add_row().cells
        p_name = row[0].paragraphs[0]
        aplicar_estilo_calibri(p_name.add_run(pers["Nombre"]), 9, pers["EsPaciente"])

        for idx, k_locus in enumerate(["A", "B", "C", "DQ", "DRB1"], 1):
            p_cell = row[idx].paragraphs[0]
            p_cell.alignment = WD_ALIGN_PARAGRAPH.CENTER
            alelos = pers["Loci"].get(k_locus, [])

            for i, alelo_obj in enumerate(alelos):
                if i > 0: p_cell.add_run("\n")
                texto_limpio = alelo_obj["val"].replace("(", "").replace(")", "")
                run_alelo = p_cell.add_run(texto_limpio)

                # APLICAR ESTILO SEGÚN EL ESTADO
                color = RGBColor(255, 0, 0) if alelo_obj.get("es_incierto") else None
                aplicar_estilo_calibri(run_alelo, 9, negrita=alelo_obj.get("bold", False), color_rgb=color)

    # --- COMENTARIOS CON FORMATO ESPECIAL ---
    doc.add_paragraph()
    p_com_h = doc.add_paragraph()
    aplicar_estilo_calibri(p_com_h.add_run("COMENTARIO:"), T_TEXTO, True)

    for nom_donante, scr in resumen.items():
        match = re.search(r'\d+', scr)
        score_val = int(match.group()) if match else 0

        p_com = doc.add_paragraph(style='List Bullet')
        p_com.paragraph_format.space_after = Pt(4)

        # Inicio de la frase con el nombre del paciente en negrita
        aplicar_estilo_calibri(p_com.add_run("El donante "), T_TEXTO)
        aplicar_estilo_calibri(p_com.add_run(f" {nom_donante} "), T_TEXTO, True)
        aplicar_estilo_calibri(p_com.add_run(f" presenta "), T_TEXTO)

        # Lógica de estados en MAYÚSCULAS y NEGRITA
        if score_val == 10:
            aplicar_estilo_calibri(p_com.add_run("compatibilidad "), T_TEXTO)
            aplicar_estilo_calibri(p_com.add_run("IDÉNTICA"), T_TEXTO, True)
            texto_score = " (10/10)."
        elif score_val == 5:
            aplicar_estilo_calibri(p_com.add_run("compatibilidad "), T_TEXTO)
            aplicar_estilo_calibri(p_com.add_run("HAPLOIDÉNTICA"), T_TEXTO, True)
            texto_score = " (5/10)."
        elif score_val == 0:
            aplicar_estilo_calibri(p_com.add_run("NO COMPATIBILIDAD"), T_TEXTO, True)
            texto_score = " (0/10)."
        else:
            aplicar_estilo_calibri(p_com.add_run("compatibilidad "), T_TEXTO)
            aplicar_estilo_calibri(p_com.add_run(f"{score_val}/10"), T_TEXTO, True)
            texto_score = "."

        aplicar_estilo_calibri(p_com.add_run(texto_score), T_TEXTO)

    doc.save(ruta_salida)

def generar_reporte(carpeta_destino):
    entrada = os.path.join(carpeta_destino, "Analisis_HLA.docx")
    salida = os.path.join(carpeta_destino, "Reporte_HLA.docx")
    try:
        datos, scores = extraer_datos_completos(entrada)
        if datos:
            generar_documento(datos, scores, salida)
            return True
        return False
    except Exception as e:
        print(f"Error en reporte final: {e}")
        return False
