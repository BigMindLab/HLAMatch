import pandas as pd
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import os


# --- FUNCIONES DE NORMALIZACIÓN Y LIMPIEZA ---

def normalizar_codigo(texto):
    if pd.isna(texto) or texto == "": return ""
    texto_str = str(texto).strip()
    if "*" in texto_str: return texto_str.split("*")[1]
    return texto_str


def limpiar_detalle(texto_detalle):
    if pd.isna(texto_detalle) or texto_detalle == "": return set()
    texto = str(texto_detalle).replace(":=:", "").strip()
    return set([normalizar_codigo(e) for e in texto.split("/") if e.strip()])


def obtener_grupo_base(texto_alelo):
    if pd.isna(texto_alelo) or texto_alelo == "": return ""
    return texto_alelo.split(":")[0]


def es_grupo_generico(texto_alelo):
    if pd.isna(texto_alelo): return False
    return "XX" in str(texto_alelo).upper()


def preparar_alelo_dinamico(fila, nombre_col_grupo):
    if nombre_col_grupo not in fila:
        return {"grupo": "", "detalle": "", "existe": False}
    grupo = fila[nombre_col_grupo]
    detalle = fila.get(f"{nombre_col_grupo}_Detalle", "")
    return {"grupo": grupo, "detalle": detalle, "existe": True}


def calcular_match_single(p_dat, d_dat):
    res = {
        "ok": False,
        "score": 0,
        "match_con": "-",
        "estado": "Sin Datos",
        "info_detalle": "",
        "n_iguales": 0,
        "n_dif": 0
    }

    if not p_dat["grupo"] or not d_dat["grupo"]:
        return res

    g_p = obtener_grupo_base(p_dat["grupo"])
    g_d = obtener_grupo_base(d_dat["grupo"])

    # --- Verificación de Grupos Base ---
    if g_p != g_d:
        res.update({
            "match_con": d_dat["grupo"],
            "estado": "No compatible",
            "info_detalle": f"Grupos distintos ({g_p} vs {g_d})"
        })
        return res

    set_p = limpiar_detalle(p_dat["detalle"])
    set_d = limpiar_detalle(d_dat["detalle"])

    # Fallback alta vs baja resolución
    if not set_p and ":" in p_dat["grupo"] and not es_grupo_generico(p_dat["grupo"]):
        set_p = {normalizar_codigo(p_dat["grupo"])}
    if not set_d and ":" in d_dat["grupo"] and not es_grupo_generico(d_dat["grupo"]):
        set_d = {normalizar_codigo(d_dat["grupo"])}

    inter = set_p & set_d
    dif = (set_p | set_d) - inter

    n_total = len(inter) + len(dif)
    res["n_iguales"] = len(inter)
    res["n_dif"] = len(dif)

    porc_iguales = len(inter) / n_total if n_total > 0 else 0

    # --- REGLA A: <40% iguales ---
    if n_total > 0 and porc_iguales < 0.4:

        detalle_base = (
            f"Iguales: {', '.join(inter)}"
            + (f"\nDiferentes: {', '.join(dif)}" if dif else "")
        )

        # Caso: hay intersección → antes era Incierto
        if len(inter) >= 1:

            # --- NUEVA REGLA: Inclusión ≥90% ---
            set_small, set_large = (
                (set_p, set_d) if len(set_p) <= len(set_d) else (set_d, set_p)
            )

            incluidos = set_small & set_large
            no_incluidos = set_small - set_large

            cobertura = len(incluidos) / len(set_small) if len(set_small) > 0 else 0

            if cobertura >= 0.9:
                res.update({
                    "ok": True,
                    "score": 1,
                    "match_con": d_dat["grupo"],
                    "estado": "Incluido",
                    "n_iguales": len(incluidos),
                    "n_dif": len(no_incluidos),
                    "info_detalle": (
                        f"Incluidos: {len(incluidos)} / {len(set_small)} ({cobertura:.0%})"
                        + (f"\nIncluidos: {', '.join(incluidos)}" if incluidos else "")
                        + (f"\nNo incluidos: {', '.join(no_incluidos)}" if no_incluidos else "")
                    )
                })
            else:
                res.update({
                    "ok": False,
                    "score": 0,
                    "match_con": d_dat["grupo"],
                    "estado": "Incierto",
                    "info_detalle": detalle_base
                })

        else:
            # 0 intersecciones → No compatible
            res.update({
                "ok": False,
                "score": 0,
                "match_con": d_dat["grupo"],
                "estado": "No compatible",
                "info_detalle": detalle_base
            })

    # --- REGLA B: Idéntico ---
    elif len(dif) == 0 and len(inter) > 0:
        res.update({
            "ok": True,
            "score": 10,
            "match_con": d_dat["grupo"],
            "estado": "Idéntico",
            "info_detalle": f"Coincidencia total: {', '.join(inter)}"
        })

    # --- REGLA C: Compatible normal ---
    else:
        detalle = (
            f"Iguales: {', '.join(inter)}"
            + (f"\nDiferentes: {', '.join(dif)}" if dif else "")
        )
        res.update({
            "ok": True,
            "score": 10 if inter else 2,
            "match_con": d_dat["grupo"],
            "estado": "Compatible",
            "info_detalle": detalle if inter else "Faltan detalles de alta resolución"
        })

    return res

def mapear_columnas_dq(df):
    if len(df) == 0: return None, None, None, None
    paciente = df.iloc[0]
    columnas = df.columns.tolist()
    candidatos = [c for c in columnas if c.startswith("DQ_") and "Detalle" not in c]
    cols_alpha = [];
    cols_beta = []
    for col in candidatos:
        val = str(paciente[col]).upper()
        if "DQA" in val:
            cols_alpha.append(col)
        elif "DQB" in val:
            cols_beta.append(col)
    return (cols_alpha[0] if len(cols_alpha) > 0 else None, cols_alpha[1] if len(cols_alpha) > 1 else None,
            cols_beta[0] if len(cols_beta) > 0 else None, cols_beta[1] if len(cols_beta) > 1 else None)


def analizar_pares_dq(paciente, donante, cols_dq):
    col_a1, col_a2, col_b1, col_b2 = cols_dq
    if not col_a1 or not col_b1: return [], 0
    p_alphas = [preparar_alelo_dinamico(paciente, col_a1)]
    if col_a2 and paciente[col_a2] != "": p_alphas.append(preparar_alelo_dinamico(paciente, col_a2))
    p_betas = [preparar_alelo_dinamico(paciente, col_b1)]
    if col_b2 and paciente[col_b2] != "": p_betas.append(preparar_alelo_dinamico(paciente, col_b2))
    d_alphas = [preparar_alelo_dinamico(donante, col_a1)]
    if col_a2 and donante[col_a2] != "": d_alphas.append(preparar_alelo_dinamico(donante, col_a2))
    d_betas = [preparar_alelo_dinamico(donante, col_b1)]
    if col_b2 and donante[col_b2] != "": d_betas.append(preparar_alelo_dinamico(donante, col_b2))

    def check_pair(h_pac, h_don):
        m_alpha = calcular_match_single(h_pac["alpha"], h_don["alpha"])
        m_beta = calcular_match_single(h_pac["beta"], h_don["beta"])

        txt_res = f"({m_alpha['match_con']} + {m_beta['match_con']})"
        detalle = f"A: {m_alpha['info_detalle']}\nB: {m_beta['info_detalle']}"

        n_ig = m_alpha.get("n_iguales", 0) + m_beta.get("n_iguales", 0)
        n_df = m_alpha.get("n_dif", 0) + m_beta.get("n_dif", 0)
        n_total = n_ig + n_df
        porc_ig = n_ig / n_total if n_total > 0 else 0

        # --- NUEVA REGLA DQ GLOBAL ---
        if n_total > 0 and porc_ig < 0.4 and n_ig >= 1:
            return {
                "score": 0,
                "estado": "Incierto",
                "match_con": txt_res,
                "info_detalle": detalle,
                "n_iguales": n_ig,
                "n_dif": n_df
            }

        # Ambos OK → Compatible
        if m_alpha["ok"] and m_beta["ok"]:
            return {
                "score": 10,
                "estado": "Compatible",
                "match_con": txt_res,
                "info_detalle": detalle,
                "n_iguales": n_ig,
                "n_dif": n_df
            }

        # Si alguno es Incierto → Incierto
        if m_alpha["estado"] == "Incierto" or m_beta["estado"] == "Incierto":
            return {
                "score": 0,
                "estado": "Incierto",
                "match_con": txt_res,
                "info_detalle": detalle,
                "n_iguales": n_ig,
                "n_dif": n_df
            }

        # Caso realmente No compatible
        return {
            "score": 0,
            "estado": "No compatible",
            "match_con": txt_res,
            "info_detalle": detalle,
            "n_iguales": n_ig,
            "n_dif": n_df
        }

    def generar_moleculas(alphas, betas):
        configs = [[{"alpha": alphas[0], "beta": betas[0]}]]
        if len(alphas) > 1 and len(betas) > 1:
            configs[0].append({"alpha": alphas[1], "beta": betas[1]})
            configs.append([{"alpha": alphas[0], "beta": betas[1]}, {"alpha": alphas[1], "beta": betas[0]}])
        return configs

    p_configs = generar_moleculas(p_alphas, p_betas)
    d_configs = generar_moleculas(d_alphas, d_betas)
    mejor_res = [];
    max_puntos = -1
    for pares_p in p_configs:
        for pares_d in d_configs:
            res_A = [];
            score_A = 0
            if len(pares_p) > 0 and len(pares_d) > 0:
                r1 = check_pair(pares_p[0], pares_d[0]);
                res_A.append((pares_p[0], r1));
                score_A += r1["score"]
                if len(pares_p) > 1 and len(pares_d) > 1:
                    r2 = check_pair(pares_p[1], pares_d[1]);
                    res_A.append((pares_p[1], r2));
                    score_A += r2["score"]
            res_B = [];
            score_B = -1
            if len(pares_p) > 0 and len(pares_d) > 1:
                score_B = 0;
                r1x = check_pair(pares_p[0], pares_d[1]);
                res_B.append((pares_p[0], r1x));
                score_B += r1x["score"]
                if len(pares_p) > 1:
                    r2x = check_pair(pares_p[1], pares_d[0]);
                    res_B.append((pares_p[1], r2x));
                    score_B += r2x["score"]
            local = res_A if score_A >= score_B else res_B
            s_local = score_A if score_A >= score_B else score_B
            if s_local > max_puntos: max_puntos = s_local; mejor_res = local

    final_res = [];
    puntos = 0
    for item in mejor_res:
        obj_p, res = item
        if res["score"] >= 10: puntos += 1
        final_res.append({"locus": "DQ Pair", "paciente_alelo": f"{obj_p['alpha']['grupo']} + {obj_p['beta']['grupo']}",
                          "resultado": res})
    return final_res, puntos


# --- ANÁLISIS GENERAL ---

def analizar_compatibilidad(df):
    if len(df) < 2: return None
    paciente = df.iloc[0];
    donantes = df.iloc[1:]
    cols_totales = df.columns.tolist()
    loci_indiv = [l for l in ["A", "B", "C", "DR", "DRB1"] if f"{l}_1" in cols_totales or l in cols_totales]
    cols_dq = mapear_columnas_dq(df)
    resultados_totales = {};
    resumen = []

    for idx, donante in donantes.iterrows():
        nombre = donante.get("Sample ID", f"D_{idx}")
        reporte = [];
        puntos_g = 0

        for locus in loci_indiv:
            p1 = preparar_alelo_dinamico(paciente, f"{locus}_1")
            p2 = preparar_alelo_dinamico(paciente, f"{locus}_2")
            if p2["grupo"] == "" and p1["existe"]: p2 = p1
            d1 = preparar_alelo_dinamico(donante, f"{locus}_1")
            d2 = preparar_alelo_dinamico(donante, f"{locus}_2")
            if d2["grupo"] == "" and d1["existe"]: d2 = d1

            mA1 = calcular_match_single(p1, d1);
            mA2 = calcular_match_single(p2, d2)
            mB1 = calcular_match_single(p1, d2);
            mB2 = calcular_match_single(p2, d1)

            sel = [(p1, mA1), (p2, mA2)] if (mA1["score"] + mA2["score"]) >= (mB1["score"] + mB2["score"]) else [
                (p1, mB1), (p2, mB2)]

            for item in sel:
                orig, res = item
                if res["estado"] == "Incluido":
                    puntos_g += 1
                elif res["score"] >= 10:
                    puntos_g += 1

                if orig["existe"] or (orig == p1):
                    reporte.append({"locus": locus, "paciente_alelo": orig["grupo"], "resultado": res})

        if cols_dq[0] and cols_dq[2]:
            res_dq, pts_dq = analizar_pares_dq(paciente, donante, cols_dq)
            reporte.extend(res_dq)
            puntos_g += pts_dq
        else:
            reporte.append({"locus": "DQ", "paciente_alelo": "Error",
                            "resultado": {"match_con": "-", "estado": "No detectado DQA/DQB", "score": 0,
                                          "info_detalle": "", "n_iguales": 0, "n_dif": 0}})

        resultados_totales[nombre] = reporte
        resumen.append({"donante": nombre, "puntaje": puntos_g,
                        "asteriscos": sum(1 for r in reporte if "Posiblemente" in r["resultado"]["estado"])})

    return paciente.get("Sample ID", "Pac"), resultados_totales, resumen


# --- GENERACIÓN DE INFORME ---

from collections import defaultdict
from docx.shared import Inches, RGBColor, Pt


def generar_word(paciente_id, resultados, resumen, nombre_salida):
    doc = Document()
    doc.add_heading(f'Analisis HLA', 0).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Paciente: {paciente_id}")

    # --- TABLA DE RESUMEN ---
    t = doc.add_table(rows=1, cols=2)
    t.style = 'Table Grid'
    t.rows[0].cells[0].text = "Donante"
    t.rows[0].cells[1].text = "Compatibilidad Total"
    for x in sorted(resumen, key=lambda i: i['puntaje'], reverse=True):
        r = t.add_row().cells
        r[0].text = str(x["donante"])
        r[1].text = f"{int(x['puntaje'])}/10"

    # --- DETALLE POR DONANTE ---
    for donante, data in resultados.items():
        doc.add_heading(f"Donante: {donante}", level=2)

        tab = doc.add_table(rows=1, cols=6)
        tab.style = 'Table Grid'
        tab.autofit = False
        anchos = [0.6, 1.1, 1.1, 1.0, 2.3, 0.4]

        h = tab.rows[0].cells
        titulos = ["Locus", "Paciente", "Donante", "Estado", "Detalles", ""]
        for i, titulo in enumerate(titulos):
            h[i].text = titulo
            h[i].width = Pt(anchos[i] * 72)

        locus_nombres_unicos = []
        for d in data:
            nombre = "DQ" if "DQ" in d["locus"] else d["locus"]
            if nombre not in locus_nombres_unicos:
                locus_nombres_unicos.append(nombre)

        for nombre_locus in locus_nombres_unicos:
            filas_grupo = [d for d in data if ("DQ" if "DQ" in d["locus"] else d["locus"]) == nombre_locus]

            # --- NUEVA LÓGICA ÚLTIMA COLUMNA ---
            estados_grupo = [f["resultado"]["estado"] for f in filas_grupo]

            n_identico = estados_grupo.count("Idéntico")
            n_compatible = estados_grupo.count("Compatible") + estados_grupo.count("Incluido")
            n_no_compatible = sum(
                1 for e in estados_grupo if e in ["No compatible", "Incierto"]
            )

            # Definición de etiquetas según tu requerimiento
            if n_no_compatible == 0:
                # Casos: (I+I), (I+C), (C+C)
                etiqueta, color = "C", RGBColor(0, 128, 0)  # Verde
            elif n_no_compatible == 1:
                # Casos: (I+NC), (C+NC)
                etiqueta, color = "H", RGBColor(204, 170, 0)  # Dorado/Amarillo
            else:
                # Caso: (NC+NC)
                etiqueta, color = "NC", RGBColor(255, 0, 0)  # Rojo

            celdas_loci = []

            for d in filas_grupo:
                row_cells = tab.add_row().cells
                for idx, w in enumerate(anchos):
                    row_cells[idx].width = Pt(w * 72)

                row_cells[0].text = d["locus"]
                row_cells[1].text = str(d["paciente_alelo"])
                res = d["resultado"]
                row_cells[2].text = str(res.get("match_con", "-"))

                estado_actual = res.get("estado", "")
                row_cells[3].text = estado_actual

                # --- LÓGICA DE DETALLES (PUNTO 1) ---
                conteo_texto = f"Iguales: {res.get('n_iguales', 0)}, Diferentes: {res.get('n_dif', 0)}"
                estado_actual = res.get("estado", "")
                det_orig = res.get("info_detalle", "")

                if estado_actual == "Idéntico":
                    # Solo ponemos el conteo
                    row_cells[4].text = conteo_texto

                elif estado_actual == "No compatible" and "Grupos distintos" in det_orig:
                    # Caso 1: Error de Grupo (Mantenemos el mensaje de texto sin la lista de alelos)
                    row_cells[4].text = f"{det_orig}\n{conteo_texto}"

                else:
                    # Mantenemos detalle original + conteo

                    row_cells[4].text = f"{det_orig}\n{conteo_texto}" if det_orig else conteo_texto

                for p in row_cells[4].paragraphs:
                    for run in p.runs: run.font.size = Pt(8)

                # Color del estado
                # --- LÓGICA DE COLOR DE ESTADO SEGURA ---
                p_estado = row_cells[3].paragraphs[0]
                p_estado.clear()  # Limpiamos para asegurar que controlamos el run
                run_st = p_estado.add_run(estado_actual)
                run_st.bold = True  # Opcional: añade negrita para que resalte más

                if estado_actual == "No compatible":
                    run_st.font.color.rgb = RGBColor(114, 47, 55)
                elif estado_actual == "Idéntico":
                    run_st.font.color.rgb = RGBColor(0, 100, 0)
                elif estado_actual == "Compatible":
                    run_st.font.color.rgb = RGBColor(60, 179, 113)
                elif estado_actual == "Incierto":
                    run_st.font.color.rgb = RGBColor(237, 33, 0)
                elif estado_actual == "Incluido":
                    run_st.font.color.rgb = RGBColor(30, 144, 255)

                celdas_loci.append(row_cells[5])

            if len(celdas_loci) > 1:
                celda_final = celdas_loci[0].merge(celdas_loci[-1])
            else:
                celda_final = celdas_loci[0]

            celda_final.text = etiqueta
            celda_final.vertical_alignment = 1
            p = celda_final.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.runs[0]
            run.bold = True
            run.font.color.rgb = color

    doc.save(nombre_salida)

def generar_informe(carpeta_destino):
    archivo_csv = os.path.join(carpeta_destino, "resultados_hla_final_limpio.csv")
    if not os.path.exists(archivo_csv):
        print(f"Error: No se encuentra {archivo_csv}")
        return False
    try:
        df = pd.read_csv(archivo_csv).fillna("")
        datos = analizar_compatibilidad(df)
        if datos:
            ruta_word = os.path.join(carpeta_destino, "Analisis_HLA.docx")
            generar_word(*datos, nombre_salida=ruta_word)
            return True
        return False
    except Exception as e:
        print(f"Error procesando informe: {e}")
        return False